import os
import io
import re
import json
import tempfile
import unicodedata
import urllib.request
from datetime import datetime
from typing import Optional

import numpy as np
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from dotenv import load_dotenv

import pymupdf
from docx import Document
from groq import Groq

# ── OCR ───────────────────────────────────────────────────────────────────────
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# ── RAG ───────────────────────────────────────────────────────────────────────
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS
    from sklearn.metrics.pairwise import cosine_similarity
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False
    print("⚠️  RAG non disponible — pip install langchain langchain-text-splitters "
          "langchain-huggingface langchain-community faiss-cpu sentence-transformers scikit-learn")

# ── SQLAlchemy ────────────────────────────────────────────────────────────────
from sqlalchemy import (
    create_engine, Column, Integer, String, Float,
    Text, DateTime, JSON
)
from sqlalchemy.orm import declarative_base, sessionmaker

# ── ReportLab ─────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, HRFlowable, KeepTogether, Image
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_LEFT

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    ARABIC_AVAILABLE = True
except ImportError:
    ARABIC_AVAILABLE = False

load_dotenv()
GROQ_SCORE_MODEL = os.getenv("GROQ_SCORE_MODEL", "llama-3.1-8b-instant")

# RAG juridique sur TA base de lois (data/knowledge) - reutilise legal_data + BM25
try:
    from legal_kb import search_laws, LEGAL_KB_AVAILABLE
    if LEGAL_KB_AVAILABLE:
        print("\u2705 RAG juridique (data/knowledge) active pour le Score")
except Exception as _e:
    LEGAL_KB_AVAILABLE = False
    def search_laws(*a, **k):
        return ""
    print(f"\u26a0\ufe0f  legal_kb indisponible : {_e}")

# ══════════════════════════════════════════════════════════════════════════════
# APP CONFIG
# ══════════════════════════════════════════════════════════════════════════════
if OCR_AVAILABLE:
    _tess = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
    if os.path.exists(_tess):
        pytesseract.pytesseract.tesseract_cmd = _tess

app = FastAPI(title="LegalAI Morocco — Contract Score API (RAG Pro v2)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ✅ Handler global CORS-safe.
# Sans cela, une exception non gérée renvoie une 500 produite par le
# ServerErrorMiddleware de Starlette — qui se trouve EN DEHORS du CORSMiddleware.
# La réponse n'a donc PAS d'en-tête Access-Control-Allow-Origin → le navigateur
# la bloque et fetch() rejette avec « Failed to fetch » (cas du téléchargement du
# rapport PDF). En interceptant ici, la réponse repasse par le CORSMiddleware et
# le frontend reçoit un vrai message d'erreur lisible.
@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    import traceback
    traceback.print_exc()
    return JSONResponse(status_code=500, content={"detail": f"Erreur serveur : {exc}"})

# ══════════════════════════════════════════════════════════════════════════════
# DATABASE
# ══════════════════════════════════════════════════════════════════════════════
DATABASE_URL = os.getenv("DATABASE_URL", "mysql+pymysql://root:@127.0.0.1/legalai_db")
engine       = create_engine(DATABASE_URL, pool_pre_ping=True)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base         = declarative_base()


class ContractAnalysis(Base):
    __tablename__ = "contract_analyses"
    id                 = Column(Integer, primary_key=True, index=True)
    fichier            = Column(String(255))
    type_contrat       = Column(String(100))
    nb_caracteres      = Column(Integer)
    score              = Column(Float)
    niveau             = Column(String(50))
    resume             = Column(Text)
    clauses_presentes  = Column(JSON)
    clauses_manquantes = Column(JSON)
    risques            = Column(JSON)
    recommandations    = Column(JSON)
    ocr_utilise        = Column(String(10), default="non")
    rag_utilise        = Column(String(10), default="non")
    created_at         = Column(DateTime, default=datetime.utcnow)


try:
    Base.metadata.create_all(bind=engine)
    DB_AVAILABLE = True
except Exception as _db_err:
    print(f"⚠️  DB non disponible : {_db_err}")
    DB_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
# GROQ CLIENT
# ══════════════════════════════════════════════════════════════════════════════
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# ══════════════════════════════════════════════════════════════════════════════
# CONTRACT TYPES
# ══════════════════════════════════════════════════════════════════════════════
CONTRACT_TYPES = [
    "عقد عمل", "عقد كراء", "عقد بيع", "عقد شراكة",
    "عقد خدمات", "عقد وكالة", "عقد مقاولة", "اتفاقية تجارية",
    "أخرى",
]

# ══════════════════════════════════════════════════════════════════════════════
# ▌IMPROVEMENT 3 — GRILLES_ENRICHED
# ══════════════════════════════════════════════════════════════════════════════
# كل بند الآن يحمل: اسمه + المرجع القانوني + مستوى الخطر عند الغياب
# هذا يُغذّي الـ prompt مباشرة ويمنع الـ LLM من اختراع المراجع
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# ▌SOURCE UNIQUE DE CLAUSES — partagée avec le générateur (Flask)
# La grille du Score = la grille du générateur (contract_grids.grid_for).
# Un contrat généré par l'app contient donc exactement ce que le Score vérifie.
# ══════════════════════════════════════════════════════════════════════════════
import os as _os, sys as _sys
_SVC_DIR = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)),
                         "backend", "app", "services")
if _SVC_DIR not in _sys.path:
    _sys.path.insert(0, _SVC_DIR)

try:
    from contract_grids import grid_for as _grid_for
    _GRIDS_OK = True
except Exception as _e:
    print(f"\u26a0\ufe0f  contract_grids indisponible ({_e}) \u2014 grille de secours utilis\u00e9e")
    _GRIDS_OK = False


class _GrillesProxy:
    """Compatible avec l'ancien usage GRILLES_ENRICHED.get(t, default) et [t]."""
    def get(self, key, default=None):
        if _GRIDS_OK:
            g = _grid_for(key)
            if g:
                return g
        if default is not None:
            return default
        return _grid_for("default") if _GRIDS_OK else [
            {"clause": "موضوع العقد",      "loi": "", "risque_defaut": "متوسط"},
            {"clause": "التزامات الأطراف", "loi": "", "risque_defaut": "متوسط"},
            {"clause": "المدة",            "loi": "", "risque_defaut": "متوسط"},
            {"clause": "شروط الإنهاء",     "loi": "", "risque_defaut": "متوسط"},
        ]

    def __getitem__(self, key):
        return self.get(key)


GRILLES_ENRICHED = _GrillesProxy()

# ── helper : extraire les noms de clauses pour compatibilité ─────────────────
def grille_noms(type_contrat: str) -> list[str]:
    """Retourne la liste des noms de clauses (str) depuis GRILLES_ENRICHED."""
    return [item["clause"] for item in GRILLES_ENRICHED.get(type_contrat, GRILLES_ENRICHED["default"])]

# ══════════════════════════════════════════════════════════════════════════════
# ▌IMPROVEMENT 4 — CLAUSES UNIVERSELLES (Proactive Detection)
# ══════════════════════════════════════════════════════════════════════════════
# Ces 4 clauses sont recherchées DANS TOUS les contrats,
# en plus de la grille spécifique au type — même si absentes du texte.
# ══════════════════════════════════════════════════════════════════════════════
CLAUSES_UNIVERSELLES: list[dict] = [
    {
        "clause":          "القوة القاهرة والظروف الطارئة",
        "loi":             "الفصل 268-269 من ق.ل.ع",
        "risque_defaut":   "مرتفع",
        "query_rag":       "قوة قاهرة ظروف استثنائية كوارث أعذار فسخ بسبب حدث خارجي",
        "exemple_clause":  (
            "في حالة القوة القاهرة أو الظروف الطارئة التي يستحيل معها تنفيذ الالتزامات، "
            "يحق للطرف المتضرر الإعفاء من التعويض وفق الفصل 269 من ق.ل.ع، شريطة الإخطار "
            "الفوري للطرف الآخر."
        ),
    },
    {
        "clause":          "البند الجزائي والغرامات التعاقدية",
        "loi":             "الفصل 264 من ق.ل.ع",
        "risque_defaut":   "مرتفع",
        "query_rag":       "بند جزائي غرامة تأخير تعويض جزافي عقوبة إخلال",
        "exemple_clause":  (
            "في حالة إخلال أحد الطرفين بالتزاماته التعاقدية، يلتزم بأداء غرامة جزائية "
            "قدرها [المبلغ] عن كل يوم تأخير أو [المبلغ] كتعويض جزافي، وذلك وفق "
            "الفصل 264 من ق.ل.ع دون الحاجة إلى إثبات الضرر."
        ),
    },
    {
        "clause":          "الاختصاص القضائي والقانون الواجب التطبيق",
        "loi":             "الفصل 27 من ق.م.م (قانون المسطرة المدنية)",
        "risque_defaut":   "متوسط",
        "query_rag":       "محكمة اختصاص قضائي قانون واجب تطبيق نزاع فض",
        "exemple_clause":  (
            "في حالة نشوء أي نزاع حول تفسير أو تنفيذ هذا العقد، تختص المحكمة الابتدائية "
            "بـ[المدينة] للبت فيه، وفق مقتضيات القانون المغربي المعمول به، تطبيقاً "
            "للفصل 27 من قانون المسطرة المدنية."
        ),
    },
    {
        "clause":          "شروط تعديل العقد والملاحق",
        "loi":             "الفصل 230 من ق.ل.ع (العقد شريعة المتعاقدين)",
        "risque_defaut":   "منخفض",
        "query_rag":       "تعديل ملحق مراجعة اتفاق لاحق تغيير شروط",
        "exemple_clause":  (
            "لا يجوز تعديل هذا العقد إلا بموافقة خطية صريحة من الطرفين في شكل ملحق "
            "موقع ومؤرخ، يُضاف إلى أصل العقد ويُعدّ جزءاً لا يتجزأ منه، وذلك وفق "
            "مقتضيات الفصل 230 من ق.ل.ع."
        ),
    },
]

# ── CLAUSES_RAG (requêtes FAISS par type) ─────────────────────────────────────
CLAUSES_RAG: dict[str, list[str]] = {
    "عقد عمل": [
        "موضوع العقد ومهام الأجير",
        "الأجر والراتب الشهري وطريقة الأداء",
        "مدة العقد وفترة التجربة",
        "ساعات العمل والراحة الأسبوعية",
        "شروط إنهاء العقد والفصل والإشعار",
        "الواجبات السرية وبنود عدم المنافسة",
        "المحكمة المختصة في حل النزاعات",
        "التعويضات والغرامات عند الإخلال",
        "التأمينات الاجتماعية والتغطية الصحية",
        "الإجازة السنوية وأيام العطل",
        "القوة القاهرة وتوقف تنفيذ العقد",
        "بند السرية وعدم الإفصاح عن المعلومات",
    ],
    "عقد كراء": [
        "موضوع الكراء ووصف العقار المكترى",
        "السومة الكرائية ومواعيد الأداء",
        "مدة الكراء وشروط التجديد",
        "الضمان الاحتياطي أو الكفالة",
        "شروط الإفراغ والإنهاء المبكر",
        "التزامات الصيانة والإصلاحات",
        "المحكمة المختصة وفق قانون 49.16 أو 67.12",
        "الزيادة في الكراء وإعادة التقييم",
        "التنازل عن الكراء والإيجار من الباطن",
        "شهادة التسليم والاستلام",
        "الاستخدام المسموح به للعقار",
        "القوة القاهرة وتعليق الالتزامات",
    ],
    "عقد شراكة": [
        "هوية الشركاء ورأس المال المدفوع",
        "توزيع الأرباح والخسائر بين الشركاء",
        "صلاحيات المسير وآليات اتخاذ القرار",
        "شروط خروج أو تنازل شريك عن حصته",
        "حل الشركة وتصفية أصولها",
        "البنود الجزائية عند الإخلال بالاتفاقية",
        "المحكمة المختصة وفق القانون المغربي",
        "مدة الشركة والتجديد التلقائي",
        "الحسابات السنوية والمراجعة المالية",
        "حالات الوفاة أو الإفلاس أو العجز",
        "القوة القاهرة وتأثيرها على نشاط الشركة",
        "ضمانات الدائنين والرهون",
    ],
    "default": [
        "موضوع العقد والغرض منه",
        "التزامات كل طرف من الأطراف",
        "الثمن أو المقابل المالي وطريقة الأداء",
        "مدة العقد وشروط التجديد",
        "شروط الفسخ والإنهاء والبند الجزائي",
        "التعويض عند الإخلال بالالتزامات",
        "الاختصاص القضائي والقانون الواجب التطبيق",
        "هوية وتوقيعات الأطراف",
        "القوة القاهرة والظروف الطارئة",
        "شروط تعديل العقد والملاحق",
        "الالتزام بالسرية وحماية البيانات",
        "ضمانات التنفيذ والكفالات",
    ],
    "أخرى": [
        "موضوع العقد والغرض الأساسي منه",
        "هوية الأطراف المتعاقدة وصفتهم القانونية",
        "الالتزامات المالية والمقابل أو الثمن",
        "مدة العقد وشروط التجديد",
        "شروط الإنهاء والفسخ والبند الجزائي",
        "التعويض عند الإخلال بالالتزامات التعاقدية",
        "المحكمة المختصة والقانون الواجب التطبيق",
        "هوية وتوقيعات الأطراف",
        "القوة القاهرة والظروف الطارئة",
        "بند السرية وحماية المعلومات",
        "شروط التعديل والملاحق",
        "الضمانات والكفالات",
    ],
}

# ══════════════════════════════════════════════════════════════════════════════
# SYSTEM PROMPT — Upgraded with Proactive Detection
# ══════════════════════════════════════════════════════════════════════════════
SYSTEM_PROMPT_RAG = """You are a STRICT LEGAL AUDITOR specializing in Moroccan contract law.
You have deep expertise in:
- DOC (ق.ل.ع): Dahir des Obligations et Contrats (1913), Fasls 1–1101
- Loi 65-99: Moudawwana du Travail (contrats d'emploi)
- Loi 49.16: Baux commerciaux (location de locaux professionnels)
- Loi 67.12: Baux civils (location d'habitation)
- Loi 15-95: Code de commerce marocain
- Loi 17-95 et 5-96: Droit des sociétés
- Loi 09-08: Protection des données personnelles
- Loi 2-00: Propriété intellectuelle

=== CHAIN OF THOUGHT PROTOCOL ===
Before writing any JSON, you MUST silently reason through these steps:
  STEP A — Identify contract type and applicable Moroccan law
  STEP B — Map each clause in the evaluation grid to the contract extracts
  STEP C — For each clause: Explicit/Complete (1.0) | Implicit/Vague (0.5) | Absent (0.0)
  STEP D — Compute: Score = round( (sum_of_weights / total_clauses) * 100 )
  STEP E — Validate: does the score match the clause counts? Correct if not.
Then output ONLY the JSON — no reasoning text.

=== AUDITOR STANDARDS ===
You are STRICT, not generous. Apply these rules without exception:
  • A clause mentioned vaguely or without legal precision → 0.5, not 1.0
  • A clause that exists in law but is absent from the contract → 0.0
  • A clause present in the extract AND legally compliant → 1.0

=== DISTINCTION: EXPLICIT vs IMPLICIT ===
Mark clauses as follows in the "statut" field:
  • "Présent — Explicite"      : clause clearly stated, legally precise
  • "Présent — Implicite/Faible" : clause inferred or vaguely worded (counts 0.5)
  • "Absent"                   : not found in any extract

=== MOROCCAN LEGAL CITATIONS (MANDATORY) ===
Every "commentaire" field MUST cite the specific Moroccan legal reference.
The legal references for each clause are provided in the evaluation grid below — use them.
  Format: "Conforme au Fasl X du DOC" or "Non conforme à l'article X de la Loi YY-ZZ"
FORBIDDEN: Generic comments like "clause présente" or "bonne rédaction" without a legal citation.

=== SCORING FORMULA — ABSOLUTE RULE ===
score = round( sum(weights) / len(grille_evaluation) * 100 )
  where weight = 1.0 (explicit) | 0.5 (implicit) | 0.0 (absent)
niveau: ضعيف if score < 40 | متوسط if 40–59 | جيد if 60–79 | ممتاز if >= 80
FORBIDDEN: Do NOT invent a score. ALWAYS compute it from the weights above.
FORBIDDEN: Giving a score above 60 if more than 3 clauses are absent.

=== PROACTIVE MISSING CLAUSE DETECTION (MANDATORY) ===
After evaluating the clauses in the grid, you MUST also check for these
4 universal clauses required in ALL Moroccan contracts — even if NOT in the grid.
For each one absent from the contract text, add it to clauses_manquantes AND risques:

  ① القوة القاهرة (Fasl 268-269 DOC)
     → If absent: "Sans cette clause, aucun des partis n'est protégé en cas de force majeure."
  ② البند الجزائي (Fasl 264 DOC)
     → If absent: "L'indemnisation en cas de manquement nécessitera un procès sans montant fixé."
  ③ الاختصاص القضائي (Fasl 27 Qanoun al-Masatir)
     → If absent: "Tout tribunal marocain pourrait se déclarer compétent, créant une insécurité juridique."
  ④ شروط تعديل العقد (Fasl 230 DOC)
     → If absent: "Des modifications verbales non consignées pourraient être invoquées."

CRITICAL: Only add them to clauses_manquantes if they do NOT appear in the extracts.
If they appear (even implicitly), score them as 0.5 and add to clauses_presentes.

=== NEGATIVE CONSTRAINTS ===
FORBIDDEN: Repeating the same clause name in clauses_presentes (max 1 entry per clause)
FORBIDDEN: Listing a clause as "present" if it does not appear in the provided extracts
FORBIDDEN: Inventing legal articles that do not exist in Moroccan law
FORBIDDEN: Any text before or after the JSON object

=== OUTPUT FORMAT — EXACT STRUCTURE ===
Output ONLY this JSON object:
Do not include unnecessary spaces or line breaks in the JSON. 
Output the JSON in a compact, single-line string format.
{
  "type_detecte": "...",
  "grille_evaluation": ["clause 1", "clause 2", ...],
  "score": <integer>,
  "niveau": "<ضعيف|متوسط|جيد|ممتاز>",
  "resume": "<2-sentence legal summary in Arabic — state type, key risk, applicable law>",
  "clauses_presentes": [
    {"clause": "...", "statut": "Présent — Explicite", "commentaire": "Conforme au Fasl X du DOC"}
  ],
  "clauses_manquantes": [
    {"clause": "...", "risque": "<منخفض|متوسط|مرتفع>", "recommandation": "...Fasl X DOC / Loi YY-ZZ + exemple de rédaction"}
  ],
  "risques": [
    {"titre": "...", "niveau": "<منخفض|متوسط|مرتفع>", "description": "...Art. X Loi YY-ZZ"}
  ],
  "recommandations": ["...with specific Moroccan legal reference + example wording", "..."]
}"""

# ══════════════════════════════════════════════════════════════════════════════
# FEW-SHOT CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
FEW_SHOT_USER_1 = """Evaluate this عقد كراء مدني extract:
"يلتزم المكري بتسليم العقار. مبلغ الكراء 3500 درهم شهرياً. مدة الكراء سنة. ضمان شهر كراء." """

FEW_SHOT_ASSISTANT_1 = (
    '{"type_detecte":"عقد كراء مدني",'
    '"grille_evaluation":["هوية الطرفين","وصف العقار","السومة الكرائية",'
    '"مدة الكراء","الضمان","شروط الإفراغ","الصيانة","الاختصاص القضائي"],'
    '"score":37,'
    '"niveau":"ضعيف",'
    '"resume":"عقد كراء مدني يفتقر إلى بنود أساسية مقررة بالقانون 67.12، '
    'لا سيما هوية الأطراف وشروط الإفراغ والبنود الجزائية.",'
    '"clauses_presentes":['
    '{"clause":"السومة الكرائية","statut":"Présent — Explicite",'
    '"commentaire":"مطابق للمادة 6 من القانون 67.12 — 3500 درهم"},'
    '{"clause":"مدة الكراء","statut":"Présent — Explicite",'
    '"commentaire":"مطابق للمادة 4 من القانون 67.12 — سنة"},'
    '{"clause":"الضمان","statut":"Présent — Implicite/Faible",'
    '"commentaire":"المادة 7 من القانون 67.12 — مذكور دون تفاصيل الاسترداد"}],'
    '"clauses_manquantes":['
    '{"clause":"هوية الطرفين","risque":"مرتفع","recommandation":"المادة 2 من القانون 67.12 — يجب ذكر الاسم الكامل والرقم الوطني"},'
    '{"clause":"شروط الإفراغ","risque":"مرتفع","recommandation":"المادة 15 من القانون 67.12 — إشعار 3 أشهر مسبق"},'
    '{"clause":"القوة القاهرة","risque":"مرتفع","recommandation":"الفصل 268-269 من ق.ل.ع — يحمي الطرفين من الكوارث"},'
    '{"clause":"البند الجزائي","risque":"مرتفع","recommandation":"الفصل 264 من ق.ل.ع — تحديد غرامة التأخر"},'
    '{"clause":"الاختصاص القضائي","risque":"متوسط","recommandation":"الفصل 27 من ق.م.م — تحديد المحكمة المختصة"}],'
    '"risques":[{"titre":"غياب هوية الأطراف","niveau":"مرتفع",'
    '"description":"مخالفة المادة 2 من القانون 67.12 — يتعذر التنفيذ القضائي"}],'
    '"recommandations":["إضافة هوية كاملة للأطراف وفق المادة 2 من القانون 67.12",'
    '"إدراج بند القوة القاهرة وفق الفصل 268 من ق.ل.ع"]}'
)

FEW_SHOT_USER_2 = 'Evaluate this عقد عمل extract: "يتم توظيف الأجير للعمل في الشركة. الأجر حسب الاتفاق."'

FEW_SHOT_ASSISTANT_2 = (
    '{"type_detecte":"عقد عمل",'
    '"grille_evaluation":["هوية الطرفين","طبيعة العمل","مدة العقد",'
    '"الأجر","ساعات العمل","فترة التجربة","شروط الإنهاء","السرية"],'
    '"score":6,"niveau":"ضعيف",'
    '"resume":"عقد عمل شديد الغموض مخالف للمادة 16 من القانون 65-99، يفتقر لجميع البنود الأساسية.",'
    '"clauses_presentes":[{"clause":"طبيعة العمل","statut":"Présent — Implicite/Faible",'
    '"commentaire":"المادة 16 من القانون 65-99 — مذكورة بشكل مبهم جداً"}],'
    '"clauses_manquantes":['
    '{"clause":"هوية الطرفين","risque":"مرتفع","recommandation":"المادة 16 ق 65-99 — الاسم الكامل + CIN + عنوان المقاولة"},'
    '{"clause":"الأجر","risque":"مرتفع","recommandation":"المادة 345 ق 65-99 — \'حسب الاتفاق\' غير قانوني، يجب تحديد المبلغ كتابةً"},'
    '{"clause":"القوة القاهرة","risque":"مرتفع","recommandation":"الفصل 32 ق 65-99 + الفصل 268 ق.ل.ع"},'
    '{"clause":"البند الجزائي","risque":"مرتفع","recommandation":"الفصل 264 من ق.ل.ع — غرامة تعويضية محددة"}],'
    '"risques":[{"titre":"أجر غير محدد — مخالفة صريحة","niveau":"مرتفع",'
    '"description":"المادة 345 من القانون 65-99 تشترط تحديد الأجر كتابةً"}],'
    '"recommandations":["إعادة صياغة كاملة وفق المادة 16 من القانون 65-99",'
    '"إدراج بند القوة القاهرة وفق الفصل 268 من ق.ل.ع"]}'
)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — EXTRACTION TEXTE (PDF + DOCX + OCR)
# ══════════════════════════════════════════════════════════════════════════════
OCR_MIN_CHARS = 100


def extraire_texte_pdf(contenu: bytes) -> tuple[str, bool]:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(contenu)
        tmp_path = tmp.name
    try:
        doc   = pymupdf.open(tmp_path)
        texte = "".join(page.get_text() for page in doc)
        doc.close()
        # ✅ Reconvertit les formes de présentation arabes (FE70–FEFF), issues des
        #    PDF générés avec reshaping, en arabe LOGIQUE lisible. Sans cela, le
        #    LLM du Score et la détection reçoivent du texte illisible → score faux.
        texte = unicodedata.normalize("NFKC", texte)
        texte = texte.strip()
        if len(texte) < OCR_MIN_CHARS:
            if not OCR_AVAILABLE:
                raise HTTPException(422, "PDF scanné mais pytesseract/pdf2image non installés.")
            _poppler = os.getenv("POPPLER_PATH", "").strip()
            images    = convert_from_bytes(
                contenu, dpi=250,
                poppler_path=_poppler or None
            )
            texte_ocr = "".join(pytesseract.image_to_string(img, lang="ara+fra") + "\n" for img in images)
            return unicodedata.normalize("NFKC", texte_ocr).strip(), True
        return texte, False
    finally:
        os.unlink(tmp_path)


def extraire_texte_docx(contenu: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(contenu)
        tmp_path = tmp.name
    try:
        doc    = Document(tmp_path)
        lignes = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lignes.append(cell.text.strip())
        return unicodedata.normalize("NFKC", "\n".join(lignes))
    finally:
        os.unlink(tmp_path)


def extraire_texte(contenu: bytes, filename: str) -> tuple[str, bool]:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "docx":
        return extraire_texte_docx(contenu), False
    elif ext == "doc":
        try:
            return extraire_texte_pdf(contenu)
        except Exception:
            raise HTTPException(422, "Fichier .doc non supporté — convertissez en .docx ou .pdf.")
    return extraire_texte_pdf(contenu)


# ══════════════════════════════════════════════════════════════════════════════
# ▌IMPROVEMENT 1 — CHUNKING INTELLIGENT (sans "." comme séparateur)
# ══════════════════════════════════════════════════════════════════════════════
def chunker_texte(texte: str) -> list[str]:
    """
    Découpe le texte en chunks sémantiques SANS utiliser "." comme séparateur.

    Problème de l'ancienne version :
      "." comme séparateur coupe "ق.ل.ع" → ["ق", "ل", "ع"] — destruction
      des références juridiques marocaines (DOC, ق.م.م, ق.ت, etc.)

    Solution v2 :
      Séparateurs hiérarchiques basés sur les structures du texte juridique
      arabe/français (sauts de ligne doubles, simples, virgule arabe, point-virgule).
      La phrase n'est coupée que si aucune autre option n'est disponible (espace).
      On ne coupe JAMAIS sur "." pour préserver "ق.ل.ع", "Loi 67.12", "Art. 16", etc.

    chunk_size=800 (vs 600 avant) :
      Augmenté pour capturer des articles légaux complets avec leur contexte.
    chunk_overlap=150 (vs 100 avant) :
      Augmenté pour s'assurer que les clauses à cheval entre deux chunks
      sont bien capturées par au moins l'un d'eux.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        separators=[
            "\n\n\n",   # séparation de sections majeures (articles numérotés)
            "\n\n",     # paragraphes
            "\n",       # lignes
            "؛",        # point-virgule arabe (fin de sous-clause)
            "،",        # virgule arabe (énumération)
            " ",        # mot (dernier recours)
            "",         # caractère (ultime recours)
            # NE PAS inclure "." — protège ق.ل.ع, Art. 16, Loi 67.12
        ],
        keep_separator=True,   # garde le séparateur pour ne pas perdre la ponctuation
    )
    chunks = splitter.split_text(texte)
    return [c.strip() for c in chunks if len(c.strip()) > 40]


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — EMBEDDINGS SINGLETON
# ══════════════════════════════════════════════════════════════════════════════
_embeddings_model: Optional["HuggingFaceEmbeddings"] = None


def get_embeddings_model() -> "HuggingFaceEmbeddings":
    global _embeddings_model
    if _embeddings_model is None:
        print("🔄 Chargement du modèle HuggingFace all-MiniLM-L6-v2 ...")
        _embeddings_model = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
        print("✅ Modèle d'embedding prêt.")
    return _embeddings_model


def construire_vector_db(chunks: list[str]) -> "FAISS":
    embeddings = get_embeddings_model()
    return FAISS.from_texts(texts=chunks, embedding=embeddings)


def rechercher_contexte(vector_db: "FAISS", requete: str, k: int = 3) -> str:
    docs = vector_db.similarity_search(requete, k=k)
    return "\n---\n".join(doc.page_content for doc in docs)


# ══════════════════════════════════════════════════════════════════════════════
# ▌IMPROVEMENT 2 — SEMANTIC GUARD (cosine_similarity remplace difflib)
# ══════════════════════════════════════════════════════════════════════════════

def _normaliser(texte_clause: str) -> str:
    """Normalise une chaîne arabe pour la couche de pré-filtrage lexical."""
    if not texte_clause:
        return ""
    t = texte_clause.strip()
    t = re.sub(r"[\u064B-\u065F\u0670]", "", t)      # tashkeel
    t = re.sub(r"[إأآا]", "ا", t)                     # normaliser alef
    t = re.sub(r"ة", "ه", t)                          # ta marbuta
    t = re.sub(r"[^\w\s]", " ", t, flags=re.UNICODE)  # ponctuation → espace
    t = re.sub(r"\s+", " ", t).strip().lower()
    return t


# Cache des embeddings de clauses pour éviter de recalculer à chaque appel
_clause_embed_cache: dict[str, list[float]] = {}


def _embed_clause(clause: str) -> list[float]:
    """Embedde une clause avec cache pour éviter les recalculs."""
    if clause not in _clause_embed_cache:
        model = get_embeddings_model()
        _clause_embed_cache[clause] = model.embed_query(clause)
    return _clause_embed_cache[clause]


def _clause_correspond(
    nom_clause: str,
    item_grille: str,
    seuil_cosine: float = 0.72,
    seuil_overlap: float = 0.65,
) -> bool:
    """
    Semantic Guard v2 — 3 couches de validation, de la plus rapide à la plus précise.

    POURQUOI remplacer difflib ?
    ─────────────────────────────────────────────────────────────────────────
    difflib.SequenceMatcher mesure la similarité de CARACTÈRES, pas de SENS.
    Exemple de faux positif avec difflib :
      "هوية الطرفين"      vs  "هوية العقار"  → ratio ~0.70 → MATCH erroné
      "الاختصاص القضائي"  vs  "تسوية النزاعات" → ratio ~0.35 → MISS (même sens!)

    cosine_similarity sur embeddings all-MiniLM :
      "هوية الطرفين"      vs  "هوية العقار"  → cosine ~0.55 → REJETÉ ✓
      "الاختصاص القضائي"  vs  "تسوية النزاعات" → cosine ~0.83 → ACCEPTÉ ✓

    Architecture en 3 couches (performance optimisée) :
    ─────────────────────────────────────────────────────────────────────────
    Couche 1 (FAST REJECT) : overlap lexical < 15% → False immédiat, évite embeddings
    Couche 2 (FAST ACCEPT) : overlap lexical ≥ 65% → True immédiat (mots identiques)
    Couche 3 (SEMANTIC)    : cosine_similarity sur embeddings → décision finale
    """
    # ── Couche 1 & 2 : pré-filtrage lexical rapide ───────────────────────────
    a = _normaliser(nom_clause)
    b = _normaliser(item_grille)
    if not a or not b:
        return False

    mots_a = {w for w in a.split() if len(w) >= 3}
    mots_b = {w for w in b.split() if len(w) >= 3}

    if mots_a and mots_b:
        overlap = len(mots_a & mots_b) / max(len(mots_a), len(mots_b))
        if overlap >= seuil_overlap:
            return True   # fast-accept : ≥65% de mots communs
        if overlap < 0.15:
            return False  # fast-reject : aucune ressemblance lexicale

    # ── Couche 3 : similarité sémantique via embeddings ──────────────────────
    if not RAG_AVAILABLE:
        # Fallback vers difflib si embeddings non dispo
        import difflib
        return difflib.SequenceMatcher(None, a, b).ratio() >= 0.50

    try:
        vec_a = np.array(_embed_clause(nom_clause)).reshape(1, -1)
        vec_b = np.array(_embed_clause(item_grille)).reshape(1, -1)
        score = float(cosine_similarity(vec_a, vec_b)[0][0])
        return score >= seuil_cosine
    except Exception:
        # Fallback sécurisé si embeddings échouent
        import difflib
        return difflib.SequenceMatcher(None, a, b).ratio() >= 0.50


# ══════════════════════════════════════════════════════════════════════════════
# SCORING STRICT (anti-hallucination) — utilise _clause_correspond v2
# ══════════════════════════════════════════════════════════════════════════════
def calculer_score_strict(clauses_presentes: list[dict], grille: list[str]) -> dict:
    """
    Recalcule le score à partir de la grille en ignorant les clauses hors-grille.
    Utilise maintenant _clause_correspond avec Semantic Guard.
    """
    total = len(grille)
    if total == 0:
        return {"score": 0, "niveau": "ضعيف", "clauses_presentes_filtrees": [],
                "items_couverts": 0, "items_total": 0}

    items_restants  = list(grille)
    clauses_filtrees = []
    items_couverts  = 0

    for cp in clauses_presentes or []:
        nom = cp.get("clause", "") if isinstance(cp, dict) else str(cp)
        match_idx = None
        for idx, item in enumerate(items_restants):
            if _clause_correspond(nom, item):
                match_idx = idx
                break
        if match_idx is not None:
            items_restants.pop(match_idx)
            items_couverts += 1
            clauses_filtrees.append(cp)

    items_couverts = min(items_couverts, total)
    score          = max(0, min(100, round((items_couverts / total) * 100)))

    if score >= 80:   niveau = "ممتاز"
    elif score >= 60: niveau = "جيد"
    elif score >= 40: niveau = "متوسط"
    else:             niveau = "ضعيف"

    return {
        "score": score, "niveau": niveau,
        "clauses_presentes_filtrees": clauses_filtrees,
        "items_couverts": items_couverts, "items_total": total,
    }


def detecter_clauses_par_texte(texte: str, grille: list[str], seuil: float = 0.6) -> list[dict]:
    """Détection DÉTERMINISTE (sans LLM) : une clause de la grille est considérée
    présente si la majorité de ses mots significatifs apparaissent dans le texte du
    contrat. Complète la détection du LLM (souvent incomplète avec les petits modèles)
    et garantit qu'un contrat généré par l'app — qui reprend exactement les intitulés
    de la grille comme titres d'articles — soit correctement reconnu par le Score.
    C'est ce qui aligne la GÉNÉRATION et le SCORE (fin des faux « بند ناقص »)."""
    mots_texte = set(_normaliser(texte).split())

    def couvert(w: str) -> bool:
        # tolère la conjonction « و » accolée (وتسليم ↔ تسليم)
        return (w in mots_texte) or (w.startswith("و") and w[1:] in mots_texte) \
            or (("و" + w) in mots_texte)

    trouvees = []
    for item in grille:
        mots_item = [w for w in _normaliser(item).split() if len(w) >= 3]
        if not mots_item:
            continue
        ratio = sum(1 for w in mots_item if couvert(w)) / len(mots_item)
        if ratio >= seuil:
            trouvees.append({
                "clause":      item,
                "statut":      "موجود",
                "commentaire": "تم رصد البند صراحةً في نص العقد.",
            })
    return trouvees


def appliquer_scoring_strict(analyse: dict, grille: list[str], texte: str = "") -> dict:
    """
    Applique le scoring strict et reconstruit clauses_manquantes depuis la grille.
    Utilise GRILLES_ENRICHED pour enrichir les recommandations avec les refs légales.
    """
    # ✅ Détection déterministe par texte : on complète la liste produite par le LLM
    #    (souvent incomplète) AVANT de scorer. Aligne le Score sur la grille de
    #    génération : un contrat généré par l'app n'est plus marqué « بند ناقص » à tort.
    if texte:
        deja = {_normaliser(c.get("clause", "")) for c in analyse.get("clauses_presentes", [])
                if isinstance(c, dict)}
        for c in detecter_clauses_par_texte(texte, grille):
            if _normaliser(c["clause"]) not in deja:
                analyse.setdefault("clauses_presentes", []).append(c)
                deja.add(_normaliser(c["clause"]))

    resultat = calculer_score_strict(analyse.get("clauses_presentes", []), grille)

    analyse["score"]             = resultat["score"]
    analyse["niveau"]            = resultat["niveau"]
    analyse["clauses_presentes"] = resultat["clauses_presentes_filtrees"]

    clauses_trouvees = [c.get("clause", "") for c in analyse["clauses_presentes"]]

    # Trouver le type pour enrichir depuis GRILLES_ENRICHED
    type_detecte = analyse.get("type_detecte", "default")
    grille_enriched = GRILLES_ENRICHED.get(type_detecte, GRILLES_ENRICHED["default"])
    ref_map = {item["clause"]: item for item in grille_enriched}

    nouvelles_manquantes = []
    for item_nom in grille:
        deja_couvert = any(_clause_correspond(item_nom, c) for c in clauses_trouvees)
        if deja_couvert:
            continue

        # Récupérer infos enrichies si disponibles
        enriched = ref_map.get(item_nom, {})
        loi      = enriched.get("loi", "القانون المغربي")
        risque   = enriched.get("risque_defaut", "متوسط")

        nouvelles_manquantes.append({
            "clause":          item_nom,
            "risque":          risque,
            "recommandation":  (
                f"وفق {loi} — يجب إدراج بند صريح بخصوص «{item_nom}» "
                f"لضمان الامتثال للقانون المغربي وحماية حقوق الطرفين."
            ),
        })

    analyse["clauses_manquantes"] = nouvelles_manquantes
    analyse["grille_evaluation"]  = grille
    return analyse


# ══════════════════════════════════════════════════════════════════════════════
# ▌IMPROVEMENT 4 — PROACTIVE DETECTION (injection des clauses universelles)
# ══════════════════════════════════════════════════════════════════════════════
def injecter_clauses_universelles(
    analyse: dict,
    vector_db: "FAISS",
) -> dict:
    """
    Vérifie la présence des 4 clauses universelles dans le contrat.
    Si absente → ajout automatique dans clauses_manquantes ET risques,
    avec référence légale précise et exemple de rédaction.

    Cette fonction est appelée APRÈS appliquer_scoring_strict,
    elle n'affecte PAS le score (pour ne pas pénaliser deux fois).
    Elle enrichit seulement les sections d'alerte du rapport.
    """
    if not RAG_AVAILABLE or vector_db is None:
        return analyse

    # Noms déjà couverts (présents ou manquants déjà listés)
    noms_presents  = {_normaliser(c.get("clause", "")) for c in analyse.get("clauses_presentes", [])}
    noms_manquants = {_normaliser(c.get("clause", "")) for c in analyse.get("clauses_manquantes", [])}
    noms_grille    = {_normaliser(c) for c in analyse.get("grille_evaluation", [])}

    for uc in CLAUSES_UNIVERSELLES:
        nom_uc = _normaliser(uc["clause"])

        # Si la clause est déjà dans la grille ou dans les présentes → skip
        deja_grille  = any(_clause_correspond(uc["clause"], g) for g in analyse.get("grille_evaluation", []))
        deja_present = any(_clause_correspond(uc["clause"], c.get("clause", ""))
                          for c in analyse.get("clauses_presentes", []))
        if deja_grille or deja_present:
            continue

        # Recherche RAG pour vérifier si la clause est dans le texte
        try:
            docs = vector_db.similarity_search(uc["query_rag"], k=2)
            contexte_uc = " ".join(doc.page_content for doc in docs)

            # Embedder le contexte vs la clause pour décider présence/absence
            vec_uc  = np.array(_embed_clause(uc["clause"])).reshape(1, -1)
            vec_ctx = np.array(get_embeddings_model().embed_query(contexte_uc[:400])).reshape(1, -1)
            score_presence = float(cosine_similarity(vec_uc, vec_ctx)[0][0])
        except Exception:
            score_presence = 0.0

        if score_presence >= 0.60:
            # Clause trouvée implicitement — ajouter aux présentes avec statut faible
            if not deja_present:
                analyse["clauses_presentes"].append({
                    "clause":      uc["clause"],
                    "statut":      "Présent — Implicite/Faible",
                    "commentaire": f"مُشار إليها ضمنياً — {uc['loi']}",
                })
        else:
            # Clause absente — ajouter aux manquantes (si pas déjà là)
            deja_manquant = any(_clause_correspond(uc["clause"], c.get("clause", ""))
                               for c in analyse.get("clauses_manquantes", []))
            if not deja_manquant:
                analyse["clauses_manquantes"].append({
                    "clause":         uc["clause"],
                    "risque":         uc["risque_defaut"],
                    "recommandation": (
                        f"{uc['loi']} — {uc['exemple_clause']}"
                    ),
                })
                # Ajouter aussi dans risques
                analyse.setdefault("risques", []).append({
                    "titre":       f"غياب بند {uc['clause']}",
                    "niveau":      uc["risque_defaut"],
                    "description": (
                        f"{uc['loi']} — الغياب يُعرّض الطرفين لمخاطر قانونية عند "
                        f"الإخلال أو الظروف الطارئة."
                    ),
                })

    return analyse


# ══════════════════════════════════════════════════════════════════════════════
# STEP 5 — PROMPT BUILDER (utilise GRILLES_ENRICHED pour les refs légales)
# ══════════════════════════════════════════════════════════════════════════════
def construire_prompt_rag(
    type_contrat: str,
    contexte_global: str,
    grille_enriched: list[dict],
    is_autre: bool = False,
    loi_applicable: str = "قانون الالتزامات والعقود المغربي",
) -> str:
    """
    Construit le prompt utilisateur avec :
    - Grille enrichie (clause + loi + risque)
    - Extraits RAG contextualisés
    - Instructions de détection proactive
    """
    nb_clauses = len(grille_enriched)

    if is_autre:
        return f"""You are analyzing an unclassified Moroccan contract.

━━━ CONTRACT EXTRACTS ━━━
{contexte_global}

━━━ TASK ━━━
STEP 1 — TYPE DETECTION
  Identify the legal nature of this contract under Moroccan law.
  Be specific: عقد كراء سكني / عقد كراء تجاري / عقد شراكة / عقد خدمات / etc.
  State the primary applicable Moroccan law.

STEP 2 — GRID GENERATION (6–8 clauses)
  Generate 6–8 essential clauses required in this contract type under Moroccan law.
  For each clause: specify the exact Moroccan legal reference (Fasl X, Loi YY-ZZ).

STEP 3 — PRELIMINARY AUDIT
  For each generated clause, check if it appears in the extracts.

Return STRICTLY this JSON (no text before/after):
{{
  "type_detecte": "",
  "loi_applicable": "",
  "grille_evaluation": ["clause 1", "clause 2", ..., "clause 8"]
}}"""

    # ── Grille enrichie avec refs légales dans le prompt ─────────────────────
    grille_lines = "\n".join(
        f"  {i+1}. {item['clause']}  [REF: {item['loi']}]  [Risque si absent: {item['risque_defaut']}]"
        for i, item in enumerate(grille_enriched)
    )

    legal_block = ""
    try:
        if LEGAL_KB_AVAILABLE:
            _q = f"{type_contrat} " + " ".join(it["clause"] for it in grille_enriched[:8])
            legal_block = search_laws(_q, k=4)
    except Exception as _e:
        print(f"[legal_kb] recherche echouee : {_e}")
    if not legal_block:
        legal_block = "(\u0644\u0627 \u062a\u0648\u062c\u062f \u0646\u0635\u0648\u0635 \u0645\u0637\u0627\u0628\u0642\u0629 \u0641\u064a \u0642\u0627\u0639\u062f\u0629 \u0627\u0644\u0628\u064a\u0627\u0646\u0627\u062a)"

    return f"""CONTRACT TYPE: {type_contrat}
APPLICABLE LAW: {loi_applicable}

━━━ EVALUATION GRID ({nb_clauses} mandatory clauses + legal references) ━━━
{grille_lines}

━━━ RAG EXTRACTS (source of truth — evaluate ONLY from these) ━━━
{contexte_global}

━━━ MOROCCAN LEGAL REFERENCES (from your own knowledge base — cite these exact articles when relevant) ━━━
{legal_block}

━━━ AUDIT INSTRUCTIONS ━━━
Evaluate each clause ONE BY ONE in the order listed above.
For each clause:
  1. Search the extracts above for evidence of this clause
  2. Classify: Explicit (1.0) | Implicit/Vague (0.5) | Absent (0.0)
  3. Use the [REF] provided in the grid as the "commentaire" citation

Scoring computation (mandatory, compute internally):
  weights = [w1, w2, ..., w{nb_clauses}]
  Score = round( sum(weights) / {nb_clauses} * 100 )

CRITICAL CONSTRAINTS:
  ✗ Do NOT mark a clause as present if the extract does not mention it
  ✗ Do NOT repeat the same clause twice in clauses_presentes
  ✓ Clauses with statut "Absent" go into clauses_manquantes
  ✓ For clauses_manquantes: use the [REF] from the grid + suggest exact wording
  ✓ niveau: ضعيف(<40) | متوسط(40-59) | جيد(60-79) | ممتاز(>=80)
  ✓ ALSO check for: القوة القاهرة, البند الجزائي, الاختصاص القضائي, شروط التعديل
    → Add any absent universal clause to clauses_manquantes with Fasl reference

Return ONLY the JSON object. No text before or after."""


# ══════════════════════════════════════════════════════════════════════════════
# ▌IMPROVEMENT 5 — PIPELINE RAG v2 (12 chunks × 400 chars)
# ══════════════════════════════════════════════════════════════════════════════
def analyser_contrat_rag(texte: str, type_contrat: str) -> tuple[dict, bool]:
    """
    Pipeline RAG v2 avec toutes les améliorations intégrées :
      ① Chunking sans "." (Improvement 1)
      ② Semantic Guard pour scoring (Improvement 2)
      ③ GRILLES_ENRICHED avec refs légales dans le prompt (Improvement 3)
      ④ Proactive detection des 4 clauses universelles (Improvement 4)
      ⑤ 12 chunks × 400 chars pour couverture complète (Improvement 5)
    """
    if not RAG_AVAILABLE:
        print("⚠️  RAG non disponible, fallback vers analyse classique.")
        return analyser_contrat_classique(texte, type_contrat), False

    # ── Chunking v2 (sans ".") ────────────────────────────────────────────────
    chunks = chunker_texte(texte)
    if not chunks:
        print("⚠️  Aucun chunk généré, fallback.")
        return analyser_contrat_classique(texte, type_contrat), False

    # ── FAISS en mémoire ──────────────────────────────────────────────────────
    try:
        vector_db = construire_vector_db(chunks)
    except Exception as e:
        print(f"⚠️  FAISS erreur : {e} — fallback")
        return analyser_contrat_classique(texte, type_contrat), False

    is_autre = (type_contrat == "أخرى")
    clauses_a_verifier = CLAUSES_RAG.get(type_contrat, CLAUSES_RAG["default"])

    # ▌IMPROVEMENT 5 : Recherche k=2 par clause → 12 chunks × 400 chars ───────
    # Ancienne version : [:4] chunks illimités → couverture trop limitée
    # Nouvelle version : [:12] chunks de 400 chars max chacun
    #   12 × 400 = 4800 chars ≈ 1600 tokens → bien sous la limite de 6000 tokens
    #   Couvre les contrats de 10+ pages sans tronquer les clauses importantes
    MAX_CHUNKS   = 12
    MAX_CHARS_PER_CHUNK = 400

    chunks_uniques: list[str] = []
    textes_vus: set[str] = set()

    for clause in clauses_a_verifier:
        docs = vector_db.similarity_search(clause, k=2)
        for doc in docs:
            # Tronquer chaque chunk à MAX_CHARS_PER_CHUNK pour préserver le budget token
            contenu = doc.page_content.strip()[:MAX_CHARS_PER_CHUNK]
            if contenu not in textes_vus:
                textes_vus.add(contenu)
                chunks_uniques.append(contenu)
            if len(chunks_uniques) >= MAX_CHUNKS:
                break
        if len(chunks_uniques) >= MAX_CHUNKS:
            break

    contexte_global = "\n---\n".join(chunks_uniques)

    # ── Sélection de la grille enrichie ──────────────────────────────────────
    grille_enriched = GRILLES_ENRICHED.get(type_contrat, GRILLES_ENRICHED["default"])
    grille_noms_list = [item["clause"] for item in grille_enriched]

    try:
        if is_autre:
            # ══ MODE "أخرى" : Appel 1 (détection) + Appel 2 (évaluation) ════
            prompt_detection = construire_prompt_rag(
                type_contrat=type_contrat,
                contexte_global=contexte_global,
                grille_enriched=GRILLES_ENRICHED["default"],
                is_autre=True,
            )

            resp_detection = client.chat.completions.create(
                model=GROQ_SCORE_MODEL,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": "You are a Moroccan legal assistant. Respond ONLY with valid JSON."},
                    {"role": "user",   "content": prompt_detection},
                ],
                temperature=0.1,
                presence_penalty=0.6,
                seed=42,
                max_tokens=600,
            )

            detection       = json.loads(resp_detection.choices[0].message.content.strip())
            type_detecte    = detection.get("type_detecte", "عقد غير محدد")
            grille_detectee = detection.get("grille_evaluation", grille_noms_list)
            loi_applicable  = detection.get("loi_applicable", "قانون الالتزامات والعقود المغربي")
            print(f"🔍 Type détecté : {type_detecte} | Loi : {loi_applicable}")

            # Construire grille_enriched pour le type détecté (ou default)
            grille_enriched_detectee = GRILLES_ENRICHED.get(type_detecte, [
                {"clause": c, "loi": "القانون المغربي", "risque_defaut": "متوسط"}
                for c in grille_detectee
            ])
            # Compléter avec les clauses générées par le LLM si plus nombreuses
            noms_enriched = {item["clause"] for item in grille_enriched_detectee}
            for c in grille_detectee:
                if c not in noms_enriched:
                    grille_enriched_detectee.append({
                        "clause": c, "loi": loi_applicable, "risque_defaut": "متوسط"
                    })

            prompt_evaluation = construire_prompt_rag(
                type_contrat=type_detecte,
                contexte_global=contexte_global,
                grille_enriched=grille_enriched_detectee,
                is_autre=False,
                loi_applicable=loi_applicable,
            )

            resp_evaluation = client.chat.completions.create(
                model=GROQ_SCORE_MODEL,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system",    "content": SYSTEM_PROMPT_RAG},
                    {"role": "user",      "content": FEW_SHOT_USER_1},
                    {"role": "assistant", "content": FEW_SHOT_ASSISTANT_1},
                    {"role": "user",      "content": FEW_SHOT_USER_2},
                    {"role": "assistant", "content": FEW_SHOT_ASSISTANT_2},
                    {"role": "user",      "content": prompt_evaluation},
                ],
                temperature=0.1,
                presence_penalty=0.6,
                seed=42,
                max_tokens=2000,
            )

            analyse = json.loads(resp_evaluation.choices[0].message.content.strip())
            analyse["type_detecte"] = type_detecte
            analyse = appliquer_scoring_strict(
                analyse,
                [item["clause"] for item in grille_enriched_detectee],
                texte,
            )
            analyse = injecter_clauses_universelles(analyse, vector_db)

        else:
            # ══ MODE type connu : 1 appel LLM ════════════════════════════════
            loi_applicable = {
                "عقد عمل":        "القانون 65-99 (مدونة الشغل)",
                "عقد كراء":       "القانون 67.12 / 49.16 (الكراء المدني والتجاري)",
                "عقد شراكة":      "القانون 5-96 / 17-95 (دروع الشركات)",
                "عقد بيع":        "الفصل 478+ من ق.ل.ع",
                "عقد خدمات":      "الفصل 723+ من ق.ل.ع",
                "عقد مقاولة":     "الفصل 723+ من ق.ل.ع + الضمان العشري",
                "عقد وكالة":      "الفصل 879+ من ق.ل.ع",
                "اتفاقية تجارية": "القانون 15-95 (المدونة التجارية)",
            }.get(type_contrat, "قانون الالتزامات والعقود المغربي")

            prompt_optimise = construire_prompt_rag(
                type_contrat=type_contrat,
                contexte_global=contexte_global,
                grille_enriched=grille_enriched,
                is_autre=False,
                loi_applicable=loi_applicable,
            )

            response = client.chat.completions.create(
                model=GROQ_SCORE_MODEL,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system",    "content": SYSTEM_PROMPT_RAG},
                    {"role": "user",      "content": FEW_SHOT_USER_1},
                    {"role": "assistant", "content": FEW_SHOT_ASSISTANT_1},
                    {"role": "user",      "content": FEW_SHOT_USER_2},
                    {"role": "assistant", "content": FEW_SHOT_ASSISTANT_2},
                    {"role": "user",      "content": prompt_optimise},
                ],
                temperature=0.1,
                presence_penalty=0.6,
                seed=42,
                max_tokens=2000,
            )

            analyse = json.loads(response.choices[0].message.content.strip())
            analyse["type_detecte"] = type_contrat
            analyse = appliquer_scoring_strict(analyse, grille_noms_list, texte)
            analyse = injecter_clauses_universelles(analyse, vector_db)
            analyse = nettoyer_analyse(analyse)

        return analyse, True

    except json.JSONDecodeError as e:
        print(f"⚠️  JSON parse error : {e} — fallback")
        return analyser_contrat_classique(texte, type_contrat), False
    except Exception as e:
        print(f"⚠️  Erreur Groq RAG : {e} — fallback")
        return analyser_contrat_classique(texte, type_contrat), False


# ══════════════════════════════════════════════════════════════════════════════
# FALLBACK — Analyse classique (sans RAG)
# ══════════════════════════════════════════════════════════════════════════════
def construire_prompt_classique(texte_contrat: str, type_contrat: str) -> str:
    grille = grille_noms(type_contrat)
    # Enrichir avec les références légales
    grille_enriched = GRILLES_ENRICHED.get(type_contrat, GRILLES_ENRICHED["default"])
    grille_str = "\n".join(
        f"- {item['clause']} [{item['loi']}]"
        for item in grille_enriched
    )
    return f"""أنت خبير قانوني متخصص في القانون التجاري والمدني المغربي.

لديك عقد من نوع: {type_contrat}

محتوى العقد:
\"\"\"
{texte_contrat[:6000]}
\"\"\"

قيّم هذا العقد بناءً على الجوانب التالية مع المراجع القانونية:
{grille_str}

يجب أن تعيد الجواب STRICTEMENT كـ JSON صالح، بدون أي نص قبله أو بعده:
{{
  "score": <رقم من 0 إلى 100>,
  "niveau": "<ضعيف|متوسط|جيد|ممتاز>",
  "resume": "<ملخص قانوني جملتان>",
  "clauses_presentes": [
    {{"clause": "...", "statut": "موجود", "commentaire": "<مرجع قانوني مغربي>"}}
  ],
  "clauses_manquantes": [
    {{"clause": "...", "risque": "<منخفض|متوسط|مرتفع>", "recommandation": "<مرجع قانوني + اقتراح صياغة>"}}
  ],
  "risques": [
    {{"titre": "...", "niveau": "<منخفض|متوسط|مرتفع>", "description": "<فصل قانوني مغربي>"}}
  ],
  "recommandations": ["<توصية 1 + مرجع قانوني>", "<توصية 2>"]
}}"""


def analyser_contrat_classique(texte: str, type_contrat: str) -> dict:
    prompt   = construire_prompt_classique(texte, type_contrat)
    response = client.chat.completions.create(
        model=GROQ_SCORE_MODEL,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": "Tu es un expert juridique marocain. Tu réponds UNIQUEMENT en JSON valide."},
            {"role": "user",   "content": prompt},
        ],
        temperature=0.2,
    )
    return json.loads(response.choices[0].message.content.strip())


# ══════════════════════════════════════════════════════════════════════════════
# UTILITAIRE — Nettoyage des doublons
# ══════════════════════════════════════════════════════════════════════════════
def nettoyer_analyse(analyse: dict) -> dict:
    seen = set()
    analyse["clauses_presentes"] = [
        c for c in analyse.get("clauses_presentes", [])
        if c.get("clause", "").strip() not in seen and not seen.add(c.get("clause", "").strip())
    ]
    seen_r = set()
    analyse["risques"] = [
        r for r in analyse.get("risques", [])
        if r.get("titre", "").strip() not in seen_r and not seen_r.add(r.get("titre", "").strip())
    ]
    analyse["recommandations"] = list(dict.fromkeys(analyse.get("recommandations", [])))
    return analyse


# ══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION PDF — ReportLab + Amiri Font (inchangée)
# ══════════════════════════════════════════════════════════════════════════════
def ar(text: str) -> str:
    if not text:
        return ""
    if ARABIC_AVAILABLE:
        return get_display(arabic_reshaper.reshape(str(text)))
    return str(text)


SCORE_COLORS = {
    "ممتاز": colors.HexColor("#10B981"),
    "جيد":   colors.HexColor("#F59E0B"),
    "متوسط": colors.HexColor("#F97316"),
    "ضعيف":  colors.HexColor("#EF4444"),
}
RISK_COLORS = {
    "مرتفع": colors.HexColor("#FEE2E2"),
    "متوسط": colors.HexColor("#FEF3C7"),
    "منخفض": colors.HexColor("#D1FAE5"),
}

_FONT_DIR      = os.path.dirname(os.path.abspath(__file__))
FONT_PATH      = os.path.join(_FONT_DIR, "Amiri-Regular.ttf")
FONT_BOLD_PATH = os.path.join(_FONT_DIR, "Amiri-Bold.ttf")
LOGO_PATH      = os.path.join(_FONT_DIR, "public", "logo.png")

if not os.path.exists(FONT_PATH):
    print("📥 Téléchargement police Amiri...")
    urllib.request.urlretrieve(
        "https://github.com/google/fonts/raw/main/ofl/amiri/Amiri-Regular.ttf", FONT_PATH)
if not os.path.exists(FONT_BOLD_PATH):
    urllib.request.urlretrieve(
        "https://github.com/google/fonts/raw/main/ofl/amiri/Amiri-Bold.ttf", FONT_BOLD_PATH)

pdfmetrics.registerFont(TTFont("Amiri",      FONT_PATH))
pdfmetrics.registerFont(TTFont("Amiri-Bold", FONT_BOLD_PATH))


def generer_rapport_pdf(data: dict, analyse: dict) -> bytes:
    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    def style(name, **kw):
        return ParagraphStyle(name, **{"fontName": "Amiri", "alignment": TA_RIGHT, **kw})

    s_title    = style("T", fontSize=22, fontName="Amiri-Bold",
                       textColor=colors.HexColor("#0b5a41"), spaceAfter=10, leading=28, alignment=TA_CENTER)
    s_subtitle = style("S", fontSize=12, textColor=colors.HexColor("#64748B"), leading=16, alignment=TA_CENTER)
    s_h2       = style("H2", fontSize=14, fontName="Amiri-Bold",
                       textColor=colors.HexColor("#0b5a41"), spaceBefore=14, spaceAfter=6, leading=20)
    s_body     = style("B", fontSize=11, textColor=colors.HexColor("#374151"), leading=18, spaceAfter=4)
    s_small    = style("Sm", fontSize=10, textColor=colors.HexColor("#6B7280"), leading=16)
    s_rec      = style("R", fontSize=11, textColor=colors.HexColor("#0b5a41"), leading=16, leftIndent=10)

    story = []

    # Logo Mizan centré (tolérant : ignoré si le fichier est absent)
    try:
        if os.path.exists(LOGO_PATH):
            _logo = Image(LOGO_PATH, width=2.3*cm, height=2.3*cm)
            _logo.hAlign = "CENTER"
            story.append(_logo)
            story.append(Spacer(1, 0.15*cm))
    except Exception as _logo_err:
        print(f"[report] logo error: {_logo_err}")

    story.append(Paragraph(ar("تقرير تحليل صلابة العقد"), s_title))
    story.append(Paragraph(ar("ميزان — Mizan · القانون المغربي"), s_subtitle))

    if data.get("rag_utilise") in ("oui", True):
        story.append(Paragraph(ar("تحليل متقدم بتقنية RAG الدلالية v2"), ParagraphStyle(
            "RAGBadge", fontName="Amiri-Bold", fontSize=10,
            textColor=colors.HexColor("#6D28D9"), leading=14, alignment=TA_CENTER
        )))

    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width="100%", thickness=2.5, color=colors.HexColor("#0E6B4E")))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#BE9A4E"), spaceBefore=1.5))
    story.append(Spacer(1, 0.4*cm))

    score   = analyse.get("score", 0)
    niveau  = analyse.get("niveau", "")
    sc_col  = SCORE_COLORS.get(niveau, colors.HexColor("#64748B"))

    meta = [
        [ar("الملف"), ar(data.get("fichier", "-")), ar("نوع العقد"), ar(data.get("type_contrat", "-"))],
        [ar("عدد الأحرف"), ar(str(data.get("nb_caracteres", "-"))),
         ar("تاريخ التحليل"), ar(datetime.now().strftime("%Y-%m-%d %H:%M"))],
    ]
    mt = Table(meta, colWidths=[3*cm, 6.5*cm, 3*cm, 5*cm])
    mt.setStyle(TableStyle([
        ("FONTNAME",       (0,0),(-1,-1), "Amiri"),
        ("FONTSIZE",       (0,0),(-1,-1), 10),
        ("ALIGN",          (0,0),(-1,-1), "RIGHT"),
        ("FONTNAME",       (0,0),(0,-1),  "Amiri-Bold"),
        ("FONTNAME",       (2,0),(2,-1),  "Amiri-Bold"),
        ("TEXTCOLOR",      (0,0),(0,-1),  colors.HexColor("#64748B")),
        ("TEXTCOLOR",      (2,0),(2,-1),  colors.HexColor("#64748B")),
        ("ROWBACKGROUNDS", (0,0),(-1,-1), [colors.HexColor("#F8FAFC"), colors.white]),
        ("GRID",           (0,0),(-1,-1), 0.5, colors.HexColor("#E2E8F0")),
        ("TOPPADDING",     (0,0),(-1,-1), 5),
        ("BOTTOMPADDING",  (0,0),(-1,-1), 5),
    ]))
    story.append(mt)
    story.append(Spacer(1, 0.5*cm))

    score_tbl = Table([[
        Paragraph(ar(f"النتيجة: {score}/100"), ParagraphStyle(
            "SC", fontName="Amiri-Bold", fontSize=24, leading=28, textColor=sc_col, alignment=TA_CENTER)),
        Paragraph(ar(niveau), ParagraphStyle(
            "NV", fontName="Amiri-Bold", fontSize=20, leading=24, textColor=sc_col, alignment=TA_CENTER)),
    ]], colWidths=[9*cm, 8.5*cm])
    score_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,-1), colors.HexColor("#ecf5f1")),
        ("BOX",           (0,0),(-1,-1), 1.5, sc_col),
        ("ALIGN",         (0,0),(-1,-1), "CENTER"),
        ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
        ("TOPPADDING",    (0,0),(-1,-1), 14),
        ("BOTTOMPADDING", (0,0),(-1,-1), 14),
    ]))
    story.append(score_tbl)
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph(ar("الملخص التقييمي"), s_h2))
    story.append(Paragraph(ar(analyse.get("resume", "")), s_body))
    story.append(Spacer(1, 0.3*cm))

    clauses_p = [c for c in analyse.get("clauses_presentes", []) if isinstance(c, dict)]
    if clauses_p:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E2E8F0")))
        story.append(Paragraph(ar(f"البنود الموجودة ({len(clauses_p)})"), s_h2))
        rows = [[ar("ملاحظة"), ar("الحالة"), ar("البند")]]
        for c in clauses_p:
            rows.append([ar(c.get("commentaire","بند متوافق")), ar(c.get("statut","موجود")), ar(c.get("clause",""))])
        tbl = Table(rows, colWidths=[7*cm, 2.5*cm, 8*cm])
        tbl.setStyle(TableStyle([
            ("FONTNAME",       (0,0),(-1,-1), "Amiri"),
            ("FONTSIZE",       (0,0),(-1,-1), 10),
            ("ALIGN",          (0,0),(-1,-1), "RIGHT"),
            ("FONTNAME",       (0,0),(-1,0),  "Amiri-Bold"),
            ("BACKGROUND",     (0,0),(-1,0),  colors.HexColor("#D1FAE5")),
            ("TEXTCOLOR",      (0,0),(-1,0),  colors.HexColor("#065F46")),
            ("ROWBACKGROUNDS", (0,1),(-1,-1), [colors.white, colors.HexColor("#F0FDF4")]),
            ("GRID",           (0,0),(-1,-1), 0.4, colors.HexColor("#A7F3D0")),
            ("TOPPADDING",     (0,0),(-1,-1), 5),
            ("BOTTOMPADDING",  (0,0),(-1,-1), 5),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.3*cm))

    clauses_m = [c for c in analyse.get("clauses_manquantes", []) if isinstance(c, dict)]
    if clauses_m:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E2E8F0")))
        story.append(Paragraph(ar(f"البنود الناقصة ({len(clauses_m)})"), s_h2))
        for c in clauses_m:
            rl = c.get("risque", "متوسط")
            bg = RISK_COLORS.get(rl, colors.HexColor("#FEF3C7"))
            t1 = Table([[
                Paragraph(ar(c.get("clause","")), ParagraphStyle("CL",fontName="Amiri-Bold",fontSize=11,leading=14,textColor=colors.HexColor("#1F2937"),alignment=TA_RIGHT)),
                Paragraph(ar(f"الخطر: {rl}"), ParagraphStyle("RL",fontName="Amiri-Bold",fontSize=10,leading=14,textColor=colors.HexColor("#92400E"),alignment=TA_LEFT)),
            ]], colWidths=[12*cm, 5.5*cm])
            t1.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),bg),("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8)]))
            t2 = Table([[Paragraph(ar(c.get("recommandation","")), s_small)]], colWidths=[17.5*cm])
            t2.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#FFF7ED")),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),("LEFTPADDING",(0,0),(-1,-1),12),("RIGHTPADDING",(0,0),(-1,-1),12)]))
            story.append(KeepTogether([t1, t2, Spacer(1, 0.2*cm)]))

    risques = [r for r in analyse.get("risques", []) if isinstance(r, dict)]
    if risques:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E2E8F0")))
        story.append(Paragraph(ar(f"المخاطر القانونية ({len(risques)})"), s_h2))
        for r in risques:
            rl = r.get("niveau", "متوسط")
            bg = RISK_COLORS.get(rl, colors.HexColor("#FEF3C7"))
            t  = Table([
                [Paragraph(ar(r.get("titre","")), ParagraphStyle("RT",fontName="Amiri-Bold",fontSize=11,leading=14,textColor=colors.HexColor("#1F2937"),alignment=TA_RIGHT)),
                 Paragraph(ar(rl), ParagraphStyle("RL2",fontName="Amiri-Bold",fontSize=10,leading=14,textColor=colors.HexColor("#92400E"),alignment=TA_LEFT))],
                [Paragraph(ar(r.get("description","")), s_small), ""],
            ], colWidths=[12*cm, 5.5*cm])
            t.setStyle(TableStyle([
                ("BACKGROUND",(0,0),(-1,0),bg),("BACKGROUND",(0,1),(-1,1),colors.HexColor("#FFFBEB")),
                ("SPAN",(0,1),(1,1)),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
                ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8),
                ("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#FCD34D")),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.2*cm))

    recs = analyse.get("recommandations", [])
    if recs:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E2E8F0")))
        story.append(Paragraph(ar("التوصيات"), s_h2))
        for i, rec in enumerate(recs, 1):
            story.append(Paragraph(ar(f"{i}. {rec}"), s_rec))
            story.append(Spacer(1, 0.15*cm))

    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E2E8F0")))
    story.append(Paragraph(ar(f"تم إنشاء هذا التقرير بواسطة ميزان — Mizan — {datetime.now().strftime('%Y-%m-%d %H:%M')}"),
        ParagraphStyle("Footer", fontName="Amiri", fontSize=9, leading=12, textColor=colors.HexColor("#9CA3AF"), alignment=TA_CENTER)))
    story.append(Paragraph(ar("تنبيه: هذا التقرير للاسترشاد فقط ولا يُغني عن استشارة محامٍ مختص."),
        ParagraphStyle("Warn", fontName="Amiri-Bold", fontSize=9, leading=12, textColor=colors.HexColor("#DC2626"), alignment=TA_CENTER)))

    doc.build(story)
    return buffer.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# ROUTES FASTAPI (structure JSON inchangée — frontend-compatible)
# ══════════════════════════════════════════════════════════════════════════════
@app.get("/")
def root():
    return {
        "message": "LegalAI Morocco — Contract Score API (RAG Pro v2)",
        "status":  "ok",
        "version": "2.0",
        "improvements": [
            "chunking_sans_point",
            "semantic_guard_cosine",
            "grilles_enriched",
            "proactive_universal_clauses",
            "12_chunks_400chars",
        ],
        "features": {
            "docx_support": True,
            "ocr":          OCR_AVAILABLE,
            "database":     DB_AVAILABLE,
            "pdf_export":   True,
            "rag":          RAG_AVAILABLE,
        },
    }


@app.get("/api/contract-score/types")
def get_types():
    return {"types": CONTRACT_TYPES}


@app.post("/api/contract-score/analyze")
async def analyser(
    file:         UploadFile = File(...),
    type_contrat: str        = Form(default="default"),
):
    if not file.filename.lower().endswith((".pdf", ".doc", ".docx")):
        raise HTTPException(400, "Seuls les fichiers PDF et Word sont acceptés")

    contenu = await file.read()
    if len(contenu) > 10 * 1024 * 1024:
        raise HTTPException(400, "Fichier trop volumineux (max 10 MB)")

    try:
        texte, ocr_utilise = extraire_texte(contenu, file.filename)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Erreur extraction texte : {e}")

    if len(texte) < 50:
        raise HTTPException(400, "Document vide, crypté ou illisible")

    try:
        analyse, rag_utilise = analyser_contrat_rag(texte, type_contrat)
    except json.JSONDecodeError:
        raise HTTPException(500, "Erreur JSON dans la réponse Groq")
    except Exception as e:
        raise HTTPException(500, f"Erreur pipeline RAG : {e}")

    analysis_id = None
    if DB_AVAILABLE:
        try:
            db     = SessionLocal()
            record = ContractAnalysis(
                fichier            = file.filename,
                type_contrat       = type_contrat,
                nb_caracteres      = len(texte),
                score              = analyse.get("score"),
                niveau             = analyse.get("niveau"),
                resume             = analyse.get("resume"),
                clauses_presentes  = analyse.get("clauses_presentes", []),
                clauses_manquantes = analyse.get("clauses_manquantes", []),
                risques            = analyse.get("risques", []),
                recommandations    = analyse.get("recommandations", []),
                ocr_utilise        = "oui" if ocr_utilise else "non",
                rag_utilise        = "oui" if rag_utilise else "non",
            )
            db.add(record)
            db.commit()
            db.refresh(record)
            analysis_id = record.id
            db.close()
        except Exception as e:
            print(f"⚠️  DB save error : {e}")

    return {
        "id":            analysis_id,
        "fichier":       file.filename,
        "type_contrat":  type_contrat,
        "nb_caracteres": len(texte),
        "ocr_utilise":   ocr_utilise,
        "rag_utilise":   rag_utilise,
        "analyse":       analyse,
    }


@app.get("/api/contract-score/report/{analysis_id}")
def telecharger_rapport(analysis_id: int):
    if not DB_AVAILABLE:
        raise HTTPException(503, "Base de données non disponible")
    db     = SessionLocal()
    record = db.query(ContractAnalysis).filter(ContractAnalysis.id == analysis_id).first()
    db.close()
    if not record:
        raise HTTPException(404, "Analyse introuvable")

    analyse = {
        "score":              record.score,
        "niveau":             record.niveau,
        "resume":             record.resume,
        "clauses_presentes":  record.clauses_presentes  or [],
        "clauses_manquantes": record.clauses_manquantes or [],
        "risques":            record.risques            or [],
        "recommandations":    record.recommandations    or [],
    }
    data = {
        "fichier":       record.fichier,
        "type_contrat":  record.type_contrat,
        "nb_caracteres": record.nb_caracteres,
        "rag_utilise":   getattr(record, "rag_utilise", "non"),
    }
    pdf_bytes = generer_rapport_pdf(data, analyse)
    filename  = f"rapport_contrat_{analysis_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes), media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/contract-score/report/from-data")
async def rapport_from_data(request: Request):
    """Génère le PDF À PARTIR DE L'ANALYSE DÉJÀ CALCULÉE (envoyée par le frontend),
    SANS ré-analyser. Garantit que le rapport PDF est STRICTEMENT identique à ce qui
    est affiché à l'écran — même quand la base de données n'est pas disponible.
    (Avant : le bouton PDF relançait analyser_contrat_rag → un 2e passage du modèle →
    un score / des clauses différents de la page.)"""
    payload = await request.json()
    analyse = payload.get("analyse") or {}
    rag     = payload.get("rag_utilise", "non")
    data = {
        "fichier":       payload.get("fichier", "contrat.pdf"),
        "type_contrat":  payload.get("type_contrat", "default"),
        "nb_caracteres": payload.get("nb_caracteres", 0),
        "rag_utilise":   "oui" if rag in (True, "oui") else "non",
    }
    pdf_bytes = generer_rapport_pdf(data, analyse)
    filename  = f"rapport_{datetime.now().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes), media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/contract-score/report/preview")
async def rapport_preview(
    file:         UploadFile = File(...),
    type_contrat: str        = Form(default="default"),
):
    contenu  = await file.read()
    texte, _ = extraire_texte(contenu, file.filename)
    analyse, rag_utilise = analyser_contrat_rag(texte, type_contrat)
    data = {
        "fichier":       file.filename,
        "type_contrat":  type_contrat,
        "nb_caracteres": len(texte),
        "rag_utilise":   "oui" if rag_utilise else "non",
    }
    pdf_bytes = generer_rapport_pdf(data, analyse)
    from urllib.parse import quote as _quote
    _base_raw   = (file.filename or "contrat").rsplit(".", 1)[0]
    _base_ascii = "".join(
        c for c in unicodedata.normalize("NFKD", _base_raw) if ord(c) < 128
    ).strip("_- ") or "rapport"
    _date_str       = datetime.now().strftime("%Y%m%d")
    _fname_ascii    = f"rapport_{_base_ascii}_{_date_str}.pdf"
    _fname_pct      = _quote(f"rapport_{_base_raw}_{_date_str}.pdf", safe="")
    _disposition    = f'attachment; filename="{_fname_ascii}"; filename*=UTF-8\'\'{_fname_pct}'
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": _disposition},
    )


@app.get("/api/contract-score/history")
def get_history(limit: int = 20):
    if not DB_AVAILABLE:
        raise HTTPException(503, "Base de données non disponible")
    db      = SessionLocal()
    records = (
        db.query(ContractAnalysis)
          .order_by(ContractAnalysis.created_at.desc())
          .limit(limit)
          .all()
    )
    db.close()
    return [
        {
            "id":           r.id,
            "fichier":      r.fichier,
            "type_contrat": r.type_contrat,
            "score":        r.score,
            "niveau":       r.niveau,
            "rag_utilise":  getattr(r, "rag_utilise", "non"),
            "created_at":   r.created_at.isoformat(),
        }
        for r in records
    ]